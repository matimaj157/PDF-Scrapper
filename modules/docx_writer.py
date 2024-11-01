from docx import Document

def find_and_replace_in_paragraphs(doc, search_text, replace_text):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if search_text in run.text:
                run.text = run.text.replace(search_text, replace_text)
    
    return doc

def write_docx(text, template, output_docx):
   doc = Document(template)

   for line in text:
        if line['style'] == 'Title':
            doc = find_and_replace_in_paragraphs(doc, "PRZEDMIOT", line['text'])
        elif line['style'] == 'New Subtitle':
            doc = find_and_replace_in_paragraphs(doc, "Nr zajęć", line['text'])
        elif line['style'] == 'Heading 1':
            doc.add_page_break()
            doc.add_paragraph(line['text'], style=line['style'])
        elif line['style'] == 'Heading 2':
            doc.add_paragraph(line['text'], style=line['style'])
            doc.add_paragraph("", style="No Spacing")
        elif line['style'] == 'Strong':
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line['text'])
            run.style = doc.styles['Strong']
        else:
            paragraph = doc.add_paragraph(line['text'])
            paragraph.style = doc.styles[line['style']]

   doc.save(output_docx)
